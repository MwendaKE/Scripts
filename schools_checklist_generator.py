'''
I created this script while I was a Math teacher at Nairobi Tumaini Highschool to help teachers
generate school checklists from Word documents by extracting pupil data, and creating organized 
summary tables with school statistics. It helped save a lot of time while doing school marketing
across Kiambu and Nairobi counties. 

'''

import os
import re
import docx
import pathlib
from docx.shared import Cm, Pt
from typing import List, Tuple

class SchoolChecklistGenerator:
    """A class to generate school checklists from Word documents"""
    
    def __init__(self):
        self.name_pattern = re.compile(r"NAME:\s*([\w\s\W]+)")
        # Combined file pattern with named groups
        self.file_pattern = re.compile(
            r"(?P<code>\d{8})\s+"
            r"(?P<school>[A-Z\s\W]+?)"
            r"(?:\s*pg\s*\d\s*-\s*\d+)?"
            r"(?:\s*\d\s*-\s*\d+)?"
            r"\.docx$",
            re.IGNORECASE
        )
    
    def find_number_of_pupils(self, doc_path: str) -> int:
        """Count the number of pupils in a Word document by searching for NAME: pattern"""
        try:
            doc = docx.Document(doc_path)
            return sum(1 for paragraph in doc.paragraphs 
                      if self.name_pattern.match(paragraph.text))
        except Exception as e:
            print(f"Error reading file {doc_path}: {e}")
            return 0
    
    def process_school_file(self, file_path: str, filename: str) -> Tuple | None:
        """Process a single school file and return its data if valid"""
        match = self.file_pattern.match(filename)
        if match:
            try:
                code = int(match.group('code'))
                school = match.group('school').strip()
                pupil_count = self.find_number_of_pupils(file_path)
                return (code, school, pupil_count, "", "", "")
            except ValueError:
                print(f"Invalid code format in file: {filename}")
        return None
    
    def search_school_files(self, directory_path: str) -> Tuple[List, List]:
        """Search through directory and process all school files"""
        valid_files = []
        invalid_files = []
        
        try:
            for filename in os.listdir(directory_path):
                file_path = os.path.join(directory_path, filename)
                
                # Skip directories and non-Word files
                if not os.path.isfile(file_path) or not filename.lower().endswith('.docx'):
                    invalid_files.append(filename)
                    continue
                
                # Process valid school files
                file_data = self.process_school_file(file_path, filename)
                if file_data:
                    valid_files.append(file_data)
                else:
                    invalid_files.append(filename)
                    
        except OSError as e:
            print(f"Error accessing directory {directory_path}: {e}")
        
        # Sort valid files by school code
        valid_files.sort(key=lambda x: x[0])
        return valid_files, invalid_files
    
    def write_to_word(self, output_path: str, records: List[Tuple], 
                     county: str, subcounty: str) -> None:
        """Generate the Word document checklist"""
        doc = docx.Document()
        
        # Add title
        doc.add_heading("SCHOOLS CHECKLIST", 0)
        
        # Add subtitle
        paragraph = doc.add_paragraph()
        subtitle = paragraph.add_run(f"{county.upper()} COUNTY, {subcounty.upper()} SUB-COUNTY")
        subtitle.underline = True
        subtitle.font.size = Pt(16)
        
        # Create table
        table = doc.add_table(rows=1, cols=7)
        table.style = "Table Grid"
        
        # Set column widths
        col_widths = [Cm(1.5), Cm(1.9), Cm(8.5), Cm(1.5), Cm(1.5), Cm(2.0), Cm(1.5)]
        for i, width in enumerate(col_widths):
            table.columns[i].width = width
        
        # Add header row
        headers = ["NO.", "CODE", "SCHOOL NAME", "PUPILS", "CHECK", "RECIPIENT", "SIGN."]
        header_cells = table.rows[0].cells
        for i, header in enumerate(headers):
            header_cells[i].text = header
            header_cells[i].paragraphs[0].runs[0].bold = True
        
        # Add data rows
        for i, record in enumerate(records, 1):
            code, school, num, rec, check, sign = record
            row_cells = table.add_row().cells
            
            row_cells[0].text = str(i)
            row_cells[1].text = str(code)
            row_cells[2].text = school
            row_cells[3].text = str(num)
            row_cells[4].text = check
            row_cells[5].text = rec
            row_cells[6].text = sign
        
        # Save document
        output_filename = os.path.join(output_path, f"{county.upper()}_{subcounty.upper()}_CHECKLIST.docx")
        doc.save(output_filename)
        print(f"Checklist saved as: {output_filename}")
    
    def run(self):
        """Main method to run the checklist generator"""
        print("SCHOOL CHECKLIST GENERATOR")
        print("--------------------------")
        
        while True:
            try:
                # Get directory input (The directory must contain various schools document files (.docx) to process)
                path_input = input("\n[*] Enter School Files Directory: ").strip()
                if not path_input:
                    continue
                    
                directory_path = pathlib.Path(path_input)
                if not directory_path.exists() or not directory_path.is_dir():
                    print("[!] Directory does not exist or is not a valid directory!")
                    continue
                
                # Get location information
                county = input("[*] Enter County (e.g. Nairobi): ").strip()
                subcounty = input("[*] Enter SubCounty (e.g. Kibra): ").strip()
                
                if not county or not subcounty:
                    print("[!] County and SubCounty are required!")
                    continue
                
                # Process files
                print("\n[+] Processing files...")
                records, invalid_files = self.search_school_files(str(directory_path))
                
                # Show invalid files if any
                if invalid_files:
                    print(f"\n[!] {len(invalid_files)} files could not be processed:")
                    for filename in invalid_files[:5]:  # Show first 5 to avoid clutter
                        print(f"    - {filename}")
                    if len(invalid_files) > 5:
                        print(f"    - ... and {len(invalid_files) - 5} more")
                
                # Generate checklist
                if records:
                    print(f"\n[+] Found {len(records)} valid school files")
                    self.write_to_word(str(directory_path), records, county, subcounty)
                    print("[+] Checklist generated successfully!")
                else:
                    print("[!] No valid school files found in the directory")
                
                # Ask to continue or exit
                continue_choice = input("\n[*] Generate another checklist? (y/n): ").strip().lower()
                if continue_choice != 'y':
                    print("Thank you for using the School Checklist Generator!")
                    break
                    
            except KeyboardInterrupt:
                print("\n\nOperation cancelled by user.")
                break
            except Exception as e:
                print(f"\n[!] An unexpected error occurred: {e}")


# Run the application
if __name__ == "__main__":
    generator = SchoolChecklistGenerator()
    generator.run()