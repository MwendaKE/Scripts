'''
Analyzes school documents for duplicate student names, code mismatches, 
and missing telephone information, generating comprehensive reports.
'''

import os
import re
import docx
from pathlib import Path
from collections import defaultdict, deque

class SchoolDocumentAnalyzer:
    """Analyzes school documents for duplicates and inconsistencies"""
    
    def __init__(self):
        self.doc_pattern = re.compile(r"(\d{8})\s*[\w\W\d\s]+\.docx$")
        self.name_pattern = re.compile(r"NAME:\s*([\w\s\W]+)")
        self.school_pattern = re.compile(r"SCHOOL:\s*(\d{8})\s*([\w\s\W]+)")
        self.tel_pattern = re.compile(r"TEL:\s*\d{10}\s*/\s*\d{10}\.")
    
    def find_document_files(self, directory_path):
        """Find all relevant DOCX files in the directory"""
        directory = Path(directory_path)
        return [f for f in directory.iterdir() 
                if f.is_file() and f.suffix.lower() == '.docx' and self.doc_pattern.match(f.name)]
    
    def extract_student_names(self, doc_path):
        """Extract all student names from a document"""
        try:
            doc = docx.Document(doc_path)
            return [match.group(1).strip() for para in doc.paragraphs 
                    if (match := self.name_pattern.match(para.text))]
        except Exception as e:
            print(f"Error reading {doc_path.name}: {e}")
            return []
    
    def extract_school_codes(self, doc_path):
        """Extract school codes from document content"""
        try:
            doc = docx.Document(doc_path)
            codes = deque()
            
            for para in doc.paragraphs:
                if match := self.school_pattern.match(para.text):
                    codes.append(match.group(1))
            
            return codes
        except Exception as e:
            print(f"Error processing school codes in {doc_path.name}: {e}")
            return deque()
    
    def check_telephone_presence(self, doc_path):
        """Check if telephone pattern exists in document"""
        try:
            doc = docx.Document(doc_path)
            for para in doc.paragraphs:
                if self.tel_pattern.search(para.text):
                    return True
            print(f"Telephone number not found in: {doc_path.name}")
            return False
        except Exception as e:
            print(f"Error checking telephone in {doc_path.name}: {e}")
            return False
    
    def find_duplicates(self, items):
        """Find duplicate items in a list"""
        count_dict = defaultdict(int)
        for item in items:
            count_dict[item] += 1
        
        return [item for item, count in count_dict.items() if count > 1]
    
    def extract_file_school_code(self, filename):
        """Extract school code from filename"""
        patterns = [
            re.compile(r"(\d{8})\s*[A-Z\s\W]+\s*pg\s*\d\s*-\s*\d+"),
            re.compile(r"(\d{8})\s*[A-Z\s\W]+\s*\d\s*-\s*\d+"),
            re.compile(r"(\d{8})")
        ]
        
        for pattern in patterns:
            if match := pattern.match(filename):
                return match.group(1)
        return None
    
    def analyze_directory(self, directory_path):
        """Analyze all documents in the directory"""
        doc_files = self.find_document_files(directory_path)
        analysis_results = []
        
        for doc_file in doc_files:
            print(f"Processing: {doc_file.name}")
            
            # Extract data
            student_names = self.extract_student_names(doc_file)
            student_duplicates = self.find_duplicates(student_names)
            school_codes = self.extract_school_codes(doc_file)
            file_school_code = self.extract_file_school_code(doc_file.name)
            has_telephone = self.check_telephone_presence(doc_file)
            
            # Check for school code mismatches
            code_mismatches = []
            for i, content_code in enumerate(school_codes, 1):
                if file_school_code and content_code != file_school_code:
                    code_mismatches.append(f"PG {i}: {file_school_code} -> {content_code}")
            
            analysis_results.append({
                'filename': doc_file.name,
                'student_count': len(student_names),
                'student_duplicates': student_duplicates,
                'code_mismatches': code_mismatches,
                'has_telephone': has_telephone
            })
        
        return analysis_results
    
    def generate_report(self, directory_path, analysis_results):
        """Generate analysis report file"""
        report_path = directory_path / f"DUPLICATE ANALYSIS - {directory_path.name.upper()}.txt"
        
        with open(report_path, 'w', encoding='utf-8') as report_file:
            report_file.write("SCHOOL DOCUMENT ANALYSIS REPORT\n")
            report_file.write("=" * 50 + "\n\n")
            
            for result in analysis_results:
                report_file.write(f"File: {result['filename']}\n")
                report_file.write("=" * (len(result['filename']) + 6) + "\n")
                report_file.write(f"Student Count: {result['student_count']}\n")
                report_file.write(f"Student Duplicates: {len(result['student_duplicates'])}\n")
                report_file.write(f"Telephone Present: {'Yes' if result['has_telephone'] else 'No'}\n")
                
                if result['student_duplicates']:
                    report_file.write("\nDuplicate Student Names:\n")
                    report_file.write("-" * 25 + "\n")
                    for duplicate in result['student_duplicates']:
                        report_file.write(f"- {duplicate}\n")
                
                if result['code_mismatches']:
                    report_file.write("\nSchool Code Mismatches:\n")
                    report_file.write("-" * 25 + "\n")
                    for mismatch in result['code_mismatches']:
                        report_file.write(f"- {mismatch}\n")
                
                report_file.write("\n" + "=" * 50 + "\n\n")
        
        return report_path


def main():
    analyzer = SchoolDocumentAnalyzer()
    
    print("School Document Analyzer")
    print("========================")
    
    while True:
        try:
            directory_path = input("\nEnter Schools Directory path: ").strip()
            directory = Path(directory_path)
            
            if not directory.exists() or not directory.is_dir():
                print("Error: Directory does not exist or is not a valid directory!")
                continue
            
            print("\nAnalyzing documents...")
            results = analyzer.analyze_directory(directory)
            
            if not results:
                print("No valid school documents found in the directory!")
                continue
            
            report_path = analyzer.generate_report(directory, results)
            print(f"\nAnalysis complete! Report saved as: {report_path}")
            
            # Ask to continue
            another = input("\nAnalyze another directory? (y/n): ").lower()
            if another != 'y':
                print("Goodbye!")
                break
                
        except KeyboardInterrupt:
            print("\nOperation cancelled.")
            break
        except Exception as e:
            print(f"An error occurred: {e}")


if __name__ == "__main__":
    main()