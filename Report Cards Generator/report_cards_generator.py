'''
This program generates report cards given three exams. Usually,
at the end of every semister/term, student are supposed to have done
at least three exams. This program will analyse the whole class data
and generate student report cards based on their position in class.

You can modify the program using KIVY or PYQT to create a powerful 
GUI APP that accepts data input directly from the user.

email: erickmwenda256@gmail.com
phone: +254 702 623 729 / +254 799 678 038
'''

import re
import docx
import time
from docx.enum.section import WD_ORIENT
from docx.shared import Pt, Cm, Mm, RGBColor, Inches
from docx.enum.text import WD_UNDERLINE

import matplotlib.pyplot as plt
 

opener_exams = [[['127'], [(52, 'B-', 8), (36, 'C-', 5), (60, 'B+', 10), (22, 'D', 3), (55, 'B', 9), (86, 'A', 12), (35, 'C-', 5), (66, 'A-', 11), (88, 'A', 12), (14, 'E', 1), (76, 'A', 12), (36, 'C-', 5)], [52, 'B-']], [['130'], [(33, 'D+', 4), (29, 'D', 3), (77, 'A', 12), (22, 'D', 3), (69, 'A-', 11), (58, 'B', 9), (47, 'C+', 7), (88, 'A', 12), (89, 'A', 12), (77, 'A', 12), (80, 'A', 12), (63, 'B+', 10)], [61, 'B+']], [['225'], [(12, 'E', 1), (36, 'C-', 5), (25, 'D', 3), (14, 'E', 1), (52, 'B-', 8), (47, 'C+', 7), (22, 'D', 3), (58, 'B', 9), (32, 'D+', 4), (55, 'B', 9), (41, 'C', 6), (42, 'C', 6)], [36, 'C-']], [['290'], [(-1, '', 0), (-1, '', 0), (-1, '', 0), (-1, '', 0), (-1, '', 0), (-1, '', 0), (-1, '', 0), (-1, '', 0), (-1, '', 0), (-1, '', 0), (-1, '', 0), (-1, '', 0)], ['', '']]]
midterm_exams = [[['127'], [(25, 'D', 3), (52, 'B-', 8), (5, 'E', 1), (96, 'A', 12), (55, 'B', 9), (96, 'A', 12), (44, 'C', 6), (52, 'B-', 8), (85, 'A', 12), (88, 'A', 12), (96, 'A', 12), (14, 'E', 1)], [59, 'B']], [['130'], [(25, 'D', 3), (11, 'E', 1), (9, 'E', 1), (25, 'D', 3), (96, 'A', 12), (77, 'A', 12), (25, 'D', 3), (69, 'A-', 11), (47, 'C+', 7), (85, 'A', 12), (90, 'A', 12), (58, 'B', 9)], [51, 'B-']], [['225'], [(22, 'D', 3), (85, 'A', 12), (22, 'D', 3), (74, 'A', 12), (22, 'D', 3), (36, 'C-', 5), (84, 'A', 12), (36, 'C-', 5), (96, 'A', 12), (47, 'C+', 7), (85, 'A', 12), (72, 'A', 12)], [56, 'B']], [['290'], [(-1, '', 0), (-1, '', 0), (-1, '', 0), (-1, '', 0), (-1, '', 0), (-1, '', 0), (-1, '', 0), (-1, '', 0), (-1, '', 0), (-1, '', 0), (-1, '', 0), (-1, '', 0)], ['', '']]]
endterm_exams = [[['127'], [(12, 'E', 1), (55, 'B', 9), (96, 'A', 12), (48, 'C+', 7), (55, 'B', 9), (80, 'A', 12), (12, 'E', 1), (55, 'B', 9), (36, 'C-', 5), (88, 'A', 12), (69, 'A-', 11), (30, 'D+', 4)], [53, 'B-']], [['130'], [(11, 'E', 1), (85, 'A', 12), (14, 'E', 1), (36, 'C-', 5), (66, 'A-', 11), (45, 'C+', 7), (36, 'C-', 5), (77, 'A', 12), (36, 'C-', 5), (65, 'A-', 11), (77, 'A', 12), (12, 'E', 1)], [46, 'C+']], [['225'], [(15, 'D-', 2), (88, 'A', 12), (63, 'B+', 10), (85, 'A', 12), (44, 'C', 6), (78, 'A', 12), (42, 'C', 6), (36, 'C-', 5), (85, 'A', 12), (44, 'C', 6), (36, 'C-', 5), (59, 'B', 9)], [56, 'B']], [['290'], [(25, 'D', 3), (36, 'C-', 5), (11, 'E', 1), (85, 'A', 12), (-1, '', 0), (80, 'A', 12), (36, 'C-', 5), (-1, '', 0), (85, 'A', 12), (-1, '', 0), (70, 'A', 12), (-1, '', 0)], [53, 'B-']]]
  
def process_exam_data(exam1, exam2, exam3):
    final_data = {}

    collection_data = list(zip(exam1, exam2, exam3))
    
    for x,y,z in collection_data:
        adm = list(zip(x[0],y[0],z[0]))[0][0]
        marks_data = list(zip(x[1],y[1],z[1]))
        results_data = list(zip(x[2],y[2],z[2]))
        
        final_data[adm] = [marks_data, results_data]
        
    return final_data
    
def calcDeviation(marks1, marks2, marks3):
    m1 = marks1[0] if marks1[0] != -1 else 0 
    m2 = marks2[0] if marks2[0] != -1 else 0
    m3 = marks3[0] if marks3[0] != -1 else 0 
    
    deviation = "-"
    
    if m1 and m2 and m3:
        deviation = int(m3 - (m1+m2)/2)
        if deviation > 0:
            deviation = f"+{deviation}"
        
    if (m2 and m3) and not m1:
        deviation = m3 - m2
        if deviation > 0:
            deviation = f"+{deviation}"
            
    if (m1 and m3) and not m2:
        deviation = m3 - m1
        if deviation > 0:
            deviation = f"+{deviation}"
            
    if (m1 and m2) and not m3:
        deviation = m2 - m1
        if deviation > 0:
            deviation = f"+{deviation}"
            
    return deviation
    
def calcDevSum(m1, m2, m3):
    deviation = "-"
    
    m1 = int(m1) if isinstance(m1, int) else 0
    m2 = int(m2) if isinstance(m2, int) else 0
    m3 = int(m3) if isinstance(m3, int) else 0
   
    if m1 and m2 and m3:
        deviation = int(m3 - (m1+m2)/2)
        if deviation > 0:
            deviation = f"+{deviation}"
        
    if (m2 and m3) and not m1:
        deviation = m3 - m2
        if deviation > 0:
            deviation = f"+{deviation}"
            
    if (m1 and m3) and not m2:
        deviation = m3 - m1
        if deviation > 0:
            deviation = f"+{deviation}"
            
    if (m1 and m2) and not m3:
        deviation = m2 - m1
        if deviation > 0:
            deviation = f"+{deviation}"
            
    return deviation
    
def color_table_text(value):
    try:
        value = int(value)
        
    except:
        value = 0
        
    if value > 0:
        return (0, 128, 0)
        
    elif value < 0:
        return (255, 0, 0)
        
    else:
        return (0, 0, 0)
    
def draw_progress_graph(x, y, annotations):
    # DOC TITLE
    plt.figure(figsize=(10, 2))
    plt.plot(x, y, linewidth=3.0, marker="o", ms=20)
    
    ax = plt.gca()
    ax.spines["top"].set_visible(False)
    ax.spines["right"].set_visible(False)
  
    for xi, yi, grade in zip(x, y, annotations):
        plt.annotate(grade,
                xy=(xi, yi), xycoords='data',
                xytext=(2, 2), textcoords='offset points',
                bbox=dict(boxstyle="round,pad=0.3", fc="yellow", alpha=0.8),
                color="blue",
                weight="bold")
    
    # Customize the plot
    plt.ylabel('Mean Score')
    plt.grid(linestyle="--", linewidth=0.2, color="b")
    plt.tight_layout()
    
    # Save the plot as an image
    plot_filename = 'plot.png'
    plt.savefig(plot_filename)
    plt.close()
        
    return plot_filename

def print_report_cards(exam1, exam2, exam3):
    filename = f"report_cards_{time.time()}.docx"
    
    studentexams = process_exam_data(exam1, exam2, exam3)
    out_of = len(studentexams)
    
    ### DOCUMENT WRITER ###
    
    doc = docx.Document()
    
    #--- SET MARGINS
    
    section = doc.sections[0]
    section.top_margin = Inches(0.3)
    section.bottom_margin = Inches(0.3)
    section.right_margin = Inches(0.8)
    section.left_margin = Inches(0.8)
    
    # ---------
    
    position = 1
    
    print(" > Generating report cards, please wait...")
  
    for adm, exams in studentexams.items():
        exam_data = exams[0]
        mean, grade = exams[1]
        
        eng1, eng2, eng3 = exam_data[0]
        eng_dev = calcDeviation(eng1, eng2, eng3)
        
        kis1, kis2, kis3 = exam_data[1]
        kis_dev = calcDeviation(kis1, kis2, kis3)
        
        mat1, mat2, mat3 = exam_data[2]
        mat_dev = calcDeviation(mat1, mat2, mat3)
        
        bio1, bio2, bio3 = exam_data[3]
        bio_dev = calcDeviation(bio1, bio2, bio3)
        
        phy1, phy2, phy3 = exam_data[4]
        phy_dev = calcDeviation(phy1, phy2, phy3)
        
        che1, che2, che3 = exam_data[5]
        che_dev = calcDeviation(che1, che2, che3)
        
        his1, his2, his3 = exam_data[6]
        his_dev = calcDeviation(his1, his2, his3)
        
        geo1, geo2, geo3 = exam_data[7]
        geo_dev = calcDeviation(geo1, geo2, geo3)
        
        cre1, cre2, cre3 = exam_data[8]
        cre_dev = calcDeviation(cre1, cre2, cre3)
        
        agr1, agr2, agr3 = exam_data[9]
        agr_dev = calcDeviation(agr1, agr2, agr3)
        
        com1, com2, com3 = exam_data[10]
        com_dev = calcDeviation(com1, com2, com3)
        
        bus1, bus2, bus3 = exam_data[11]
        bus_dev = calcDeviation(bus1, bus2, bus3)
  
        printing_data = [
               (100, "ENGLISH", f"{eng1[0]}  {eng1[1]}", f"{eng2[0]}  {eng2[1]}", f"{eng3[0]}  {eng3[1]}", eng_dev, "Good", "Ondari"),
               (101, "KISWAHILI", f"{kis1[0]}  {kis1[1]}", f"{kis2[0]}  {kis2[1]}", f"{kis3[0]}  {kis3[1]}", kis_dev, "Good", "Polly"),
               (202, "MATHEMATICS", f"{mat1[0]}  {mat1[1]}", f"{mat2[0]}  {mat2[1]}", f"{mat3[0]}  {mat3[1]}", mat_dev, "Very Good", "Partrick"),
               (123, "BIOLOGY", f"{bio1[0]}  {bio1[1]}", f"{bio2[0]}  {bio2[1]}", f"{bio3[0]}  {bio3[1]}", bio_dev, "Good", "Charles"),
               (342, "PHYSICS", f"{phy1[0]}  {phy1[1]}", f"{phy2[0]}  {phy2[1]}", f"{phy3[0]}  {phy3[1]}", phy_dev, "Excellent", "Wycliffe"),
               (333, "CHEMISTRY", f"{che1[0]}  {che1[1]}", f"{che2[0]}  {che2[1]}", f"{che3[0]}  {che3[1]}", che_dev, "Good", "Wycliffe"),
               (331, "HISTORY", f"{his1[0]}  {his1[1]}", f"{his2[0]}  {his2[1]}", f"{his3[0]}  {his3[1]}", his_dev, "Good", "Sandra"),
               (165, "GEOGRAPHY", f"{geo1[0]}  {geo1[1]}", f"{geo2[0]}  {geo2[1]}", f"{geo3[0]}  {geo3[1]}", geo_dev, "Good", "Miriam"),
               (112, "CRE", f"{cre1[0]}  {cre1[1]}", f"{cre2[0]}  {cre2[1]}", f"{cre3[0]}  {cre3[1]}", cre_dev, "Poor", "Doris"),
               (213, "AGRICULTURE", f"{agr1[0]}  {agr1[1]}", f"{agr2[0]}  {agr2[1]}", f"{agr3[0]}  {agr3[1]}", agr_dev, "Good", "Charles"),
               (445, "COMPUTER", f"{com1[0]}  {com1[1]}", f"{com2[0]}  {com2[1]}", f"{com3[0]}  {com3[1]}", com_dev, "Good", "Erick"),
               (559, "BUSINESS", f"{bus1[0]}  {bus1[1]}", f"{bus2[0]}  {bus2[1]}", f"{bus3[0]}  {bus3[1]}", bus_dev, "Exellent", "Jackline"),
            ] 
        
        dev_sum_tt = calcDevSum(mean[0],mean[1],mean[2])
            
        totals_data = [
               ("", "", "", "", "", "", ""),
               ("MEAN", "", mean[0], mean[1], mean[2], dev_sum_tt, ""),
               ("GRADE", "", grade[0], grade[1], grade[2], "", ""),
               ("", "", "", "", "", "", ""),
            ]
           
        # WRITE TITLE OF DOCUMENT
    
        title = doc.add_heading()
    
        title1 = title.add_run("NEPTUNE ACADEMY\n")
        title1.font.size = Pt(22)
        title1.font.name = "Times New Roman"
        #title1.font.color.rgb = RGBColor.from_string("0000FF")
          
        title2 = title.add_run("PRIMARY, JUNIOR AND SENIOR SCHOOLS\n")
        title2.font.size = Pt(16)
        title2.font.name = "Times New Roman"
        title2.font.color.rgb = RGBColor.from_string("FF0000")
           
        title3 = title.add_run("P.O BOX 11722 — 00100, UMOJA, NAIROBI\n")
        title3.font.size = Pt(12)
        title3.font.name = "Times New Roman"
        title3.font.color.rgb = RGBColor.from_string("000000")
           
        title_format = title.paragraph_format
        title_format.line_spacing = Pt(20)
        title_format.space_before = Pt(0)
        title_format.space_after = Pt(0)
        title.alignment = 1
         
        # ------------- #
    
        header = doc.add_heading()
        header_run = header.add_run("STUDENT REPORT CARD")
        header_run.font.name = "Impact"
        header_run.font.size = Pt(14)
        header_run.font.color.rgb = RGBColor.from_string("0047AB")
        header_run.underline = True
        
        header_format = header.paragraph_format
        header_format.space_before = Pt(0)
        header_format.space_after = Pt(3)
        header.alignment = 1
        
        ## WRITE STUDENT DETAILS ##
        
        name_para = doc.add_paragraph()
        paragraph_format = name_para.paragraph_format
        paragraph_format.line_spacing = Pt(13)

        name_para.add_run("NAME ")

        std_name = name_para.add_run(f"     FATUMA ABDI      ")
        std_name.underline = True
        std_name.underline = WD_UNDERLINE.DOTTED
        std_name.bold = True
        std_name.font.name = "Lucida Calligraphy"

        name_para.add_run("ADM NO. ")

        adm_no = name_para.add_run(f"      {adm}      ")
        adm_no.underline = True
        adm_no.underline = WD_UNDERLINE.DOTTED
        adm_no.bold = True
        adm_no.font.name = "Lucida Calligraphy"

        name_para.add_run("FORM ")

        form = name_para.add_run(f"      FORM 1     ")
        form.underline = True
        form.underline = WD_UNDERLINE.DOTTED
        form.bold = True

        name_para.add_run("TERM ")

        term = name_para.add_run(f"     TERM 2     ")
        term.underline = True
        term.underline = WD_UNDERLINE.DOTTED
        term.bold = True
        term.font.name = "Lucida Calligraphy"

        name_para.add_run("EXAM")

        exam = name_para.add_run(f"     ENDTERM     ")
        exam.underline = True
        exam.underline = WD_UNDERLINE.DOTTED
        exam.bold = True
        exam.font.name = "Lucida Calligraphy"

        # YEAR
        
        name_para.add_run("YEAR ")

        year = name_para.add_run(f"     2024     _")
        year.underline = True
        year.underline = WD_UNDERLINE.DOTTED
        year.bold = True
        year.font.name = "Lucida Calligraphy"
       
        position_para = doc.add_paragraph()
        paragraph_format = position_para.paragraph_format
        paragraph_format.line_spacing = Pt(15)
        
        # POSITION
        
        '''
        To get the actual position of student in class, you need to sort
        the studentexams data in descending order. The student with the highest marks
        will be in position number 1.
        '''
        position_para.add_run("MEAN GRADE ")

        mgrade = position_para.add_run(f"         {grade[-1]}        ")
        mgrade.underline = True
        mgrade.underline = WD_UNDERLINE.DOTTED
        mgrade.bold = True
        mgrade.font.name = "Lucida Calligraphy"
        
        position_para.add_run("POSITION")
        
        pstn = position_para.add_run(f"          {position}          ")
        pstn.underline = True
        pstn.underline = WD_UNDERLINE.DOTTED
        pstn.bold = True
        pstn.font.name = "Lucida Calligraphy"
        
        # OUT OF
        
        position_para.add_run("OUT OF")
        
        outof = position_para.add_run(f"          {out_of}         _")
        outof.underline = True
        outof.underline = WD_UNDERLINE.DOTTED
        outof.bold = True
        outof.font.name = "Lucida Calligraphy"

        # ADD MARKS 

        table1 = doc.add_table(rows = 1, cols = 8)
        table1.style = doc.styles["Table Grid"]

        table1.columns[0].width = Cm(1.5)
        table1.columns[1].width = Cm(6.5)
        table1.columns[2].width = Cm(3.5)
        table1.columns[3].width = Cm(3.5)
        table1.columns[4].width = Cm(3.5)
        table1.columns[5].width = Cm(2.5)
        table1.columns[6].width = Cm(8.0)
        table1.columns[7].width = Cm(4.5)

        cells = table1.rows[0].cells

        cells[0].text = "CODE"
        cells[0].width = Cm(1.5)

        cells[1].text = "SUBJECTS"
        cells[1].width = Cm(6.5)

        cells[2].text = "Opener"
        cells[2].width = Cm(3.5)
        
        cells[3].text = "Midterm"
        cells[3].width = Cm(3.5)
        
        cells[4].text = "Endterm"
        cells[4].width = Cm(3.5)
        
        cells[5].text = "Dev"
        cells[5].width = Cm(2.5)

        cells[6].text = "REMARKS"
        cells[6].width = Cm(8.0)

        cells[7].text = "TEACHER"
        cells[7].width = Cm(4.5)
           
        #----
        for row in printing_data:
            code, subject, marks1, marks2, marks3, deviation, remark, sign = row
            
            r, g , b = color_table_text(deviation)
            
            row_cells = table1.add_row().cells
        
            row_cells[0].text = str(code)
            row_cells[1].text = subject
            row_cells[2].text = str(marks1) if not re.match("-1",marks1) else "" 
            row_cells[3].text = str(marks2) if not re.match("-1",marks2) else "" 
            row_cells[4].text = str(marks3) if not re.match("-1",marks3) else "" 
            row_cells[5].text = str(deviation)
            row_cells[5].paragraphs[0].runs[0].font.color.rgb = RGBColor(r, g , b)
            row_cells[6].text = remark
            row_cells[7].text = sign
            
        for row in totals_data:
            rowname, _, mean1, mean2, mean3, devs, _ = row
            
            r, g , b = color_table_text(devs)
            
            row_cells = table1.add_row().cells
        
            row_cells[0].merge(row_cells[1])
        
            row_cells[0].text = rowname
            row_cells[1].text = ""
            row_cells[2].text = str(mean1)
            row_cells[3].text = str(mean2)
            row_cells[4].text = str(mean3)
            row_cells[5].text = str(devs)
            row_cells[5].paragraphs[0].runs[0].font.color.rgb = RGBColor(r, g , b)

            row_cells[6].text = ""
            row_cells[7].text = ""
        
        ####----------
        clstrs_comment = "Your progress is amazing. I have no doubt that next time you will soar as high as an eagle. Keep up the good work and never give up."
        hdtrs_comment = "You are a reader. Keep up your good work and never give up. Success is fo those who study well."
        sch_closing_date = "Monday 01 August, 2024"
        sch_opening_date= "Wednesday 27 August, 2024"
        
        progress_heading = doc.add_heading()
        remarks_ = progress_heading.add_run("GRAPHICAL PROGRESS", 0)
        remarks_.font.name = "Times New Roman"
        remarks_.font.size = Pt(11)
        remarks_.font.color.rgb = RGBColor.from_string("0047AB")
         
        prog_format = progress_heading.paragraph_format
        prog_format.space_after = Pt(1)
        prog_format.space_after = Pt(1)
        
        # Draw Progress Graph
        
        x_vals = ["Opener", "Midterm", "Endterm"]
        y_vals = [mean[0], mean[1], mean[2]]
         
        graph = draw_progress_graph(x_vals, y_vals, grade)
        
        ## Add Progress Graph To Document
        
        graph_para = doc.add_paragraph()
        graph_para_run = graph_para.add_run().add_picture(graph, width=Inches(5))
        graph_para.alignment = 1
        #doc.add_picture(graph, width=Inches(5))
        
        ####----------####
        
        ct_trs_remark = doc.add_paragraph()
        ct_beg = ct_trs_remark.add_run("CLASS TEACHER'S COMMENTS:")
        ct_beg.font.color.rgb = RGBColor.from_string("0047AB")
        ct_beg.font.size = Pt(11)
        ct_beg.bold = True
        
        ct_comment = ct_trs_remark.add_run(f"     {clstrs_comment}   _\n")
        ct_comment.underline = True
        ct_comment.underline = WD_UNDERLINE.DOTTED
        ct_comment.bold = True
        ct_comment.font.name = "Lucida Calligraphy"
        
        ct_trs_remark.add_run(f"Date: ")
        ct_trs_remark.add_run("."*50)
        ct_trs_remark.add_run(" "*10)
        ct_trs_remark.add_run("Signature: ")
        ct_trs_remark.add_run("."*50)

        next_term = doc.add_paragraph("School has closed today on ")
        term_date_closing = next_term.add_run(f"       {sch_closing_date}         _")
        term_date_closing.underline = True
        term_date_closing.underline = WD_UNDERLINE.DOTTED
        term_date_closing.bold = True
        term_date_closing.font.name = "Lucida Calligraphy"
        
        next_term.add_run(" and reopens next time on ") 
        term_date_opening = next_term.add_run(f"       {sch_opening_date}         _")
        term_date_opening.underline = True
        term_date_opening.underline = WD_UNDERLINE.DOTTED
        term_date_opening.bold = True
        term_date_opening.font.name = "Lucida Calligraphy"
        
        parent_seen = doc.add_paragraph("Parent / Guardian's Signature:  ")
        parent_seen.add_run("."*80)
     
        ##### GRADING SYSTEM ######
        
        grd = doc.add_paragraph("")
        
        data = [{'_id': 'DB_GRD_001', 'SCIENCE': {'E': [[0, 14], 1], 'D-': [[15, 19], 2], 'D': [[20, 29], 3], 'D+': [[30, 34], 4], 'C-': [[35, 39], 5], 'C': [[40, 44], 6], 'C+': [[45, 49], 7], 'B-': [[50, 54], 8], 'B': [[55, 59], 9], 'B+': [[60, 64], 10], 'A-': [[65, 69], 11], 'A': [[70, 100], 12]}, 'LANGUAGES': {'E': [[0, 14], 1], 'D-': [[15, 19], 2], 'D': [[20, 29], 3], 'D+': [[30, 34], 4], 'C-': [[35, 39], 5], 'C': [[40, 44], 6], 'C+': [[45, 49], 7], 'B-': [[50, 54], 8], 'B': [[55, 59], 9], 'B+': [[60, 64], 10], 'A-': [[65, 69], 11], 'A': [[70, 100], 12]}, 'HUMANITIES': {'E': [[0, 14], 1], 'D-': [[15, 19], 2], 'D': [[20, 29], 3], 'D+': [[30, 34], 4], 'C-': [[35, 39], 5], 'C': [[40, 44], 6], 'C+': [[45, 49], 7], 'B-': [[50, 54], 8], 'B': [[55, 59], 9], 'B+': [[60, 64], 10], 'A-': [[65, 69], 11], 'A': [[70, 100], 12]}, 'GENERAL': {'E': [[7, 10], 1, 'POOR'], 'D-': [[11, 17], 2, 'WEAK'], 'D': [[18, 24], 3, 'WEAK'], 'D+': [[25, 31], 4, 'WEAK'], 'C-': [[32, 38], 5, 'AVERAGE'], 'C': [[39, 45], 6, 'AVERAGE'], 'C+': [[46, 52], 7, 'AVERAGE'], 'B-': [[53, 59], 8, 'GOOD'], 'B': [[60, 66], 9, 'GOOD'], 'B+': [[67, 73], 10, 'GOOD'], 'A-': [[74, 80], 11, 'VERY GOOD'], 'A': [[81, 84], 12, 'VERY GOOD']}}]
        
        grading1 = data[0]["SCIENCE"]
        grading2 = data[0]["GENERAL"]

        form12_run = grd.add_run("FORM 1 & 2 (MARKS): ")
        form12_run.font.size = Pt(10)
        form12_run.underline = True
 
        for k,v in grading1.items():
            grdvals = "–".join([str(i) for i in v[0]])
            grd_key = grd.add_run(f"{k}: ")
            grd_key.font.size = Pt(10)
            grd_key.bold = True
            
            grd_val = grd.add_run(f"{grdvals};  ")
            grd_val.font.size = Pt(10)
            
        form34_run = grd.add_run("FORM 3 & 4 (POINTS): ")
        form34_run.font.size = Pt(10)
        form34_run.underline = True
        
        for k,v in grading2.items():
            grdvals = "–".join([str(i) for i in v[0]])
            grd_key = grd.add_run(f"{k}: ")
            grd_key.font.size = Pt(10)
            grd_key.bold = True
            
            grd_val = grd.add_run(f"{grdvals};  ")
            grd_val.font.size = Pt(10)
            
        #-------------
        
        doc.add_page_break()
        
        position += 1
        
    doc.save(filename)
    
    print(" > Report cards generated and saved to storage...")
    
    return
    
print_report_cards(opener_exams,midterm_exams,endterm_exams)
